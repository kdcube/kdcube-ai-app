# SPDX-License-Identifier: MIT
# Copyright (c) 2025 Elena Viter

# chat/sdk/tools/docx_renderer.py

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Dict, Optional, Tuple
import re
import json
import kdcube_ai_app.apps.chat.sdk.tools.md_utils as md_utils
from kdcube_ai_app.apps.chat.sdk.tools.citations import normalize_sources_any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement

# --------------------------- Theme dataclass ----------------------------------

@dataclass(frozen=True)
class DocxTheme:
    """Configurable theme for DOCX rendering — palette, type scale, and mono font."""

    # --- colour palette ---
    palette: Dict[str, RGBColor] = field(default_factory=lambda: {
        "fg": RGBColor(20, 24, 31),
        "muted": RGBColor(95, 106, 121),
        "accent": RGBColor(31, 111, 235),
        "quote_bg": RGBColor(245, 247, 250),
        "rule": RGBColor(220, 224, 230),
        "table_header_bg": RGBColor(240, 244, 252),
        "code_bg": RGBColor(250, 250, 252),
    })

    # --- type scale ---
    type_scale: Dict[str, Pt] = field(default_factory=lambda: {
        "title": Pt(22),
        "h1": Pt(18),
        "h2": Pt(16),
        "h3": Pt(14),
        "h4": Pt(12.5),
        "h5": Pt(12),
        "h6": Pt(11.5),
        "body": Pt(11.5),
        "code": Pt(10.5),
    })

    # --- mono font ---
    mono_font: str = "Consolas"

    # --- helpers ---
    def hex(self, key: str) -> str:
        """Return 6-char uppercase hex string for a palette colour (for OxmlElement attrs)."""
        c = self.palette[key]
        return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"


# Pre-built themes -------------------------------------------------------

DEFAULT_THEME = DocxTheme()

KDCUBE_THEME = DocxTheme(
    palette={
        "fg": RGBColor(0x0D, 0x1E, 0x2C),          # text  #0D1E2C
        "muted": RGBColor(0x7A, 0x99, 0xB0),        # text-muted  #7A99B0
        "accent": RGBColor(0x43, 0x72, 0xC3),        # blue  #4372C3
        "quote_bg": RGBColor(0xEE, 0xF8, 0xF7),      # surface-2  #EEF8F7
        "rule": RGBColor(0xD8, 0xEC, 0xEB),           # border  #D8ECEB
        "table_header_bg": RGBColor(0xDD, 0xEA, 0xFE),  # blue-pale  #DDEAFE
        "code_bg": RGBColor(0xF6, 0xFA, 0xFA),        # bg (light teal tint)
    },
)

# Legacy module-level aliases (kept for backward compat; prefer DocxTheme) ----
PALETTE = DEFAULT_THEME.palette
TYPE = DEFAULT_THEME.type_scale
MONO = DEFAULT_THEME.mono_font

# --------------------------- Helpers / constants -----------------------------

_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")  # ![alt](url)
_URL_RE = re.compile(r"https?://[^\s<>\"]+|www\.[^\s<>\"]+")  # Plain URLs
_CIT_RE  = re.compile(r"\[\[S:(\d+)\]\]")  # [[S:3]]
_CODE_FENCE_RE = re.compile(r"^```(\w+)?\s*$")
_TABLE_ROW_RE  = re.compile(r"^\s*\|.+\|\s*$")
_BLOCKQUOTE_RE = re.compile(r"^\s*>\s+(.*)$")
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")  # Detect headings from # to ######

def _basename_only(path: str, default_ext: str = ".docx") -> str:
    name = Path(path).name
    if default_ext and not name.lower().endswith(default_ext):
        name += default_ext
    return name

def _ensure_parent(p: Path) -> None:
    p.parent.mkdir(parents=True, exist_ok=True)

def _domain_of(url: str) -> str:
    from urllib.parse import urlparse
    try:
        net = urlparse(url).netloc
        return net or url
    except Exception:
        return url

def _add_char_style(run, *, size: Pt, bold=False, italic=False, color: RGBColor | None = None, mono=False, theme: DocxTheme = None):
    font = run.font
    font.size = size
    font.bold = bold
    font.italic = italic
    mono_name = (theme or DEFAULT_THEME).mono_font
    if mono:
        font.name = mono_name
        # East Asian fallback – keeps mono look on Windows
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), mono_name)
    if color is not None:
        font.color.rgb = color

def _set_para(p, *, space_before=3, space_after=3, line_spacing=1.25, align=None):
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if align:
        p.alignment = align

def _add_heading(doc: Document, text: str, level: int, *, theme: DocxTheme = None):
    """Add a heading with proper styling and indentation for levels 1-6."""
    t = theme or DEFAULT_THEME
    # Map to built-in styles where available
    if level <= 3:
        style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]
        p = doc.add_paragraph(style=style)
        p.clear()
    else:
        # For levels 4-6, create custom styled paragraphs
        p = doc.add_paragraph()
        # Indent deeper headings slightly
        if level >= 4:
            p.paragraph_format.left_indent = Inches(0.25 * (level - 3))

    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Set spacing based on level
    space_before = max(12 - level * 2, 3)
    space_after = max(8 - level * 2, 2)
    _set_para(p, space_before=space_before, space_after=space_after)

    r = p.add_run(text.strip())

    # Font sizes for different levels
    size_map = {
        1: t.type_scale["h1"],
        2: t.type_scale["h2"],
        3: t.type_scale["h3"],
        4: t.type_scale["h4"],
        5: t.type_scale["h5"],
        6: t.type_scale["h6"],
    }
    size = size_map.get(level, t.type_scale["h6"])

    # All headings are bold
    _add_char_style(r, size=size, bold=True, color=t.palette["fg"], theme=t)

def _add_paragraph_text(doc: Document, text: str, level: int = 0, *, theme: DocxTheme = None):
    t = theme or DEFAULT_THEME
    # Use Word built-ins for lists to keep bullets nice
    list_style = None
    if re.match(r"^\s*(?:[-*])\s+", text):
        list_style = "List Bullet"
        text = re.sub(r"^\s*[-*]\s+", "", text)
    elif re.match(r"^\s*\d+\.\s+", text):
        list_style = "List Number"
        text = re.sub(r"^\s*\d+\.\s+", "", text)
    p = doc.add_paragraph(style=list_style) if list_style else doc.add_paragraph()
    # indent by level (2 spaces → one level)
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.32 * level)

    _set_para(p, space_before=2, space_after=2)

    # Split for **bold** and *italic*
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            _add_char_style(r, size=t.type_scale["body"], bold=True, color=t.palette["fg"], theme=t)
        elif part.startswith("*") and part.endswith("*"):
            r = p.add_run(part[1:-1])
            _add_char_style(r, size=t.type_scale["body"], italic=True, color=t.palette["fg"], theme=t)
        else:
            _emit_link_or_text(p, part, theme=t)

    # default any unstyled runs
    for r in p.runs:
        if r.font.size is None:
            _add_char_style(r, size=t.type_scale["body"], color=t.palette["fg"], theme=t)

def _emit_link_or_text(p, text: str, sources_map: Dict[int, Dict[str, str]] = None, resolve_citations: bool = False, *, theme: DocxTheme = None):
    t = theme or DEFAULT_THEME
    sources_map = sources_map or {}
    while text:
        m_link = _LINK_RE.search(text)
        m_url = _URL_RE.search(text)
        m_cit = _CIT_RE.search(text) if resolve_citations else None

        # Build candidates list: prefer markdown links over plain URLs
        cands = [(m_link, "link"), (m_url, "url"), (m_cit, "cit")]
        cands = [(m, tp) for m, tp in cands if m]

        if not cands:
            r = p.add_run(text)
            return

        m, kind = min(cands, key=lambda x: x[0].start())

        # Add text before the match
        if m.start() > 0:
            r = p.add_run(text[:m.start()])

        if kind == "link":
            # Markdown link: [text](url)
            r = p.add_run(m.group(1))
            _add_char_style(r, size=t.type_scale["body"], color=t.palette["accent"], theme=t)
            try:
                r.hyperlink.address = m.group(2)
            except Exception:
                pass
        elif kind == "url":
            # Plain URL: https://example.com
            url = m.group(0)
            # Add www. prefix for www URLs without protocol
            if url.startswith("www."):
                hyperlink_url = "http://" + url
            else:
                hyperlink_url = url
            r = p.add_run(url)
            _add_char_style(r, size=t.type_scale["body"], color=t.palette["accent"], theme=t)
            try:
                r.hyperlink.address = hyperlink_url
            except Exception:
                pass
        else:  # citation
            sid = int(m.group(1))
            rec = sources_map.get(sid, {})
            label = rec.get("title") or f"[{sid}]"
            url = rec.get("url", "")
            r = p.add_run(label)
            _add_char_style(r, size=t.type_scale["body"], color=t.palette["accent"], theme=t)
            if url:
                try:
                    r.hyperlink.address = url
                except Exception:
                    pass
        text = text[m.end():]

def _add_image(doc: Document, path: str, alt_text: str = "", max_width: float = 6.0, *, theme: DocxTheme = None):
    """
    Add an image to the document.

    Args:
        doc: Document object
        path: Path to image file (local path or could be URL - caller should handle download)
        alt_text: Alternative text for the image
        max_width: Maximum width in inches (default 6.0 for standard page width)
        theme: Optional DocxTheme for styling
    """
    t = theme or DEFAULT_THEME
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para(p, space_before=6, space_after=6)

        # Add the image
        run = p.add_run()
        picture = run.add_picture(path, width=Inches(max_width))

        # Add caption if alt text provided
        if alt_text:
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_para(caption, space_before=2, space_after=8)
            r = caption.add_run(alt_text)
            _add_char_style(r, size=Pt(10), italic=True, color=t.palette["muted"], theme=t)

    except Exception as e:
        # Fallback: add error message if image can't be loaded
        p = doc.add_paragraph()
        r = p.add_run(f"[Image not found: {path}]")
        _add_char_style(r, size=t.type_scale["body"], italic=True, color=t.palette["muted"], theme=t)

def _add_code_block(doc: Document, lines: List[str], *, theme: DocxTheme = None):
    t = theme or DEFAULT_THEME
    # Use 1x1 table as a card w/ background + border for better reliability
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell = tbl.cell(0,0)
    # background
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), t.hex("code_bg"))
    cell._tc.get_or_add_tcPr().append(shading)
    # border
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for side in ("top","bottom","left","right"):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:color'), t.hex("rule"))
        tc_borders.append(el)
    tc_pr.append(tc_borders)

    p = cell.paragraphs[0]
    _set_para(p, space_before=3, space_after=3, line_spacing=1.1)
    for i, ln in enumerate(lines):
        if i > 0:
            p = cell.add_paragraph()
            _set_para(p, space_before=0, space_after=0, line_spacing=1.1)
        r = p.add_run(ln.rstrip("\n"))
        _add_char_style(r, size=t.type_scale["code"], mono=True, color=t.palette["fg"], theme=t)

def _add_blockquote(doc: Document, lines: List[str], *, theme: DocxTheme = None):
    t = theme or DEFAULT_THEME
    # 1x1 table for shaded quote with left rule
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0,0)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), t.hex("quote_bg"))
    cell._tc.get_or_add_tcPr().append(shading)

    # left rule via table borders
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for side in ("top","bottom","right"):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'nil')
        tc_borders.append(el)
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '16')
    left.set(qn('w:color'), t.hex("rule"))
    tc_borders.append(left)
    tc_pr.append(tc_borders)

    p = cell.paragraphs[0]
    _set_para(p, space_before=2, space_after=2)
    r = p.add_run("\n".join(lines))
    _add_char_style(r, size=t.type_scale["body"], italic=True, color=t.palette["muted"], theme=t)

def _is_separator_row(cells: List[str]) -> bool:
    """
    Check if a row is a markdown table separator.
    A separator cell contains only dashes (-) and optionally colons (:) for alignment.
    Examples: "---", ":---", "---:", ":---:", "------------"
    """
    if not cells:
        return False

    for cell in cells:
        # Remove whitespace
        cell = cell.strip()
        if not cell:
            return False
        # Check if cell contains only dashes and colons (at least 3 chars)
        if len(cell) < 3:
            return False
        # Must contain at least one dash
        if '-' not in cell:
            return False
        # Can only contain dashes and colons
        if not all(c in '-:' for c in cell):
            return False

    return True

def _parse_table(block_lines: List[str]) -> Optional[List[List[str]]]:
    """
    Parse markdown table lines into a list of rows.
    Returns None if the block doesn't form a valid table.
    """
    rows = [ln.strip() for ln in block_lines if _TABLE_ROW_RE.match(ln)]
    if len(rows) < 2:
        return None

    def split_row(r: str):
        return [c.strip() for c in r.strip("|").split("|")]

    cells = [split_row(r) for r in rows]

    # Check if second row is a separator
    if not _is_separator_row(cells[1]):
        return None

    # Extract header and data rows (skip separator)
    hdr = cells[0]
    data = cells[2:] if len(cells) > 2 else []

    # Ensure all rows have the same number of columns as header
    num_cols = len(hdr)
    normalized_data = []
    for row in data:
        # Pad or truncate to match header column count
        while len(row) < num_cols:
            row.append("")
        normalized_data.append(row[:num_cols])

    return [hdr] + normalized_data

def _add_table(doc: Document, data: List[List[str]], *, theme: DocxTheme = None):
    """
    Add a formatted table to the document.
    """
    t = theme or DEFAULT_THEME
    if not data or len(data) < 1:
        return

    rows, cols = len(data), len(data[0])
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.style = 'Table Grid'

    # Set column widths to be more evenly distributed
    for col_idx in range(cols):
        for row in tbl.rows:
            row.cells[col_idx].width = Inches(6.0 / cols)

    # Format header row
    for j, txt in enumerate(data[0]):
        cell = tbl.cell(0, j)
        p = cell.paragraphs[0]
        p.clear()
        r = p.add_run(txt)
        _add_char_style(r, size=t.type_scale["body"], bold=True, color=t.palette["fg"], theme=t)

        # Light header background
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), t.hex("table_header_bg"))
        cell._tc.get_or_add_tcPr().append(shading)

        # Center align header text
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Format data rows
    for i in range(1, rows):
        for j, txt in enumerate(data[i]):
            cell = tbl.cell(i, j)
            p = cell.paragraphs[0]
            p.clear()
            r = p.add_run(txt)
            _add_char_style(r, size=t.type_scale["body"], color=t.palette["fg"], theme=t)

            # Left align data cells
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def _split_markdown_sections(md: str) -> List[Tuple[str, List[str]]]:
    """
    Create sections by '## ' headings; first '# ' becomes doc title if present.
    Preserve all other heading levels in the body.
    """
    lines = (md or "").splitlines()
    slides: List[Tuple[str, List[str]]] = []
    cur_title: Optional[str] = None
    cur_body: List[str] = []

    for ln in lines:
        if ln.startswith("## "):
            if cur_title is not None:
                slides.append((cur_title.strip(), cur_body))
            cur_title = ln[3:]
            cur_body = []
        elif ln.startswith("# "):
            # Only use first # as document title
            if cur_title is None and not slides:
                cur_title = ln[2:]
                cur_body = []
            else:
                # Other # headings go in body
                cur_body.append(ln)
        else:
            cur_body.append(ln)

    if cur_title is None:
        nonempty = next((l for l in lines if l.strip()), "Document")
        cur_title = nonempty.lstrip("# ").strip() or "Document"

    slides.append((cur_title.strip(), cur_body))
    return slides

# --------------------------- Public entrypoint -------------------------------

def render_docx(
        path: str,
        content_md: str,
        *,
        title: Optional[str] = None,
        sources: Optional[list[dict] | dict] = None,
        resolve_citations: bool = True,
        include_sources_section: bool = True,
        theme: Optional[DocxTheme] = None,
) -> str:
    """
    Render a modern-looking .docx from Markdown with headings, lists, code, quotes, tables, links, citations, and images.

    Supported markdown features:
    - Headings: # through ###### (all levels)
    - Lists: bullet (-,*) and numbered (1.)
    - Bold: **text**
    - Italic: *text*
    - Links: [text](url) or plain URLs (https://example.com, www.example.com)
    - Images: ![alt text](path/to/image.png)
    - Code blocks: ```language
    - Tables: | header | header |
    - Blockquotes: > text
    - Citations: [[S:n]]

    Args:
        theme: Optional DocxTheme for palette/type/font customisation.
               Defaults to DEFAULT_THEME when omitted.

    Returns the **basename** written inside OUTPUT_DIR.
    """
    t = theme or DEFAULT_THEME

    basename = _basename_only(path, ".docx")
    outfile = Path(path)
    _ensure_parent(outfile)

    sources_map: Dict[int, Dict[str, str]] = {}
    order: List[int] = []
    if sources:
        norm_sources = normalize_sources_any(sources)
        sources_map, order = md_utils._normalize_sources(norm_sources)

    sections = _split_markdown_sections(content_md or "")
    doc = Document()

    # Title
    title_text = title or sections[0][0]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_para(p, space_before=0, space_after=4, line_spacing=1.2)
    r = p.add_run(title_text.strip())
    _add_char_style(r, size=t.type_scale["title"], bold=True, color=t.palette["fg"], theme=t)

    # subtle rule
    rule = doc.add_paragraph()
    _set_para(rule, space_before=0, space_after=8)
    r2 = rule.add_run("─" * 40)
    _add_char_style(r2, size=Pt(10), color=t.palette["rule"], theme=t)

    # Content (each section = H1 + body)
    for stitle, body in sections:
        # section heading
        _add_heading(doc, stitle, level=1, theme=t)

        in_code = False
        code_buf: List[str] = []
        table_buf: List[str] = []
        quote_buf: List[str] = []

        def flush_code():
            nonlocal code_buf
            if code_buf:
                _add_code_block(doc, code_buf, theme=t)
                code_buf = []

        def flush_table():
            nonlocal table_buf
            if table_buf:
                data = _parse_table(table_buf)
                if data:
                    _add_table(doc, data, theme=t)
                else:
                    # Fallback: render as paragraphs if parsing fails
                    for raw in table_buf:
                        _add_paragraph_text(doc, raw, theme=t)
                table_buf = []

        def flush_quote():
            nonlocal quote_buf
            if quote_buf:
                _add_blockquote(doc, quote_buf, theme=t)
                quote_buf = []

        for ln in body:
            # Handle code fences
            if _CODE_FENCE_RE.match(ln):
                if not in_code:
                    flush_table()
                    flush_quote()
                    in_code = True
                    code_buf = []
                else:
                    flush_code()
                    in_code = False
                continue

            if in_code:
                code_buf.append(ln)
                continue

            # Handle images (before other inline processing)
            m_img = _IMAGE_RE.match(ln.strip())
            if m_img:
                flush_code()
                flush_table()
                flush_quote()
                alt_text, img_path = m_img.groups()
                _add_image(doc, img_path, alt_text, theme=t)
                continue

            # Handle table rows
            if _TABLE_ROW_RE.match(ln):
                flush_code()
                flush_quote()
                table_buf.append(ln)
                continue
            else:
                flush_table()

            # Handle blockquotes
            m_q = _BLOCKQUOTE_RE.match(ln)
            if m_q:
                flush_code()
                quote_buf.append(m_q.group(1))
                continue
            else:
                flush_quote()

            # Handle headings (###, ####, etc.)
            m_heading = _HEADING_RE.match(ln)
            if m_heading:
                flush_code()
                hashes, heading_text = m_heading.groups()
                level = len(hashes)
                _add_heading(doc, heading_text, level=level, theme=t)
                continue

            # Paragraph / list
            # compute indent level from leading spaces
            m_bullet = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", ln)
            if m_bullet:
                spaces, _, tail = m_bullet.groups()
                level = min(len(spaces) // 2, 4)
                _add_paragraph_text(doc, ln, level=level, theme=t)
            else:
                if ln.strip():
                    _add_paragraph_text(doc, ln, theme=t)

        # Flush any remaining buffers
        flush_code()
        flush_table()
        flush_quote()

    # Add sources section if requested
    if include_sources_section and order:
        _add_heading(doc, "References", level=1, theme=t)
        for sid in order:
            src = sources_map.get(sid)
            if not src:
                continue
            p = doc.add_paragraph()
            _set_para(p, space_before=1, space_after=1)
            r1 = p.add_run(f"[{sid}] ")
            _add_char_style(r1, size=t.type_scale["body"], bold=True, color=t.palette["fg"], theme=t)
            r2 = p.add_run(src.get("title") or _domain_of(src.get("url","")) or f"Source {sid}")
            _add_char_style(r2, size=t.type_scale["body"], color=t.palette["fg"], theme=t)
            url = src.get("url","")
            if url:
                p2 = doc.add_paragraph()
                _set_para(p2, space_before=0, space_after=6)
                r3 = p2.add_run(url)
                _add_char_style(r3, size=t.type_scale["body"], color=t.palette["accent"], theme=t)

    doc.save(str(outfile))
    return basename


# Test with sample markdown including deep headings
if __name__ == "__main__":
    test_md = """
# Pacific Northwest Ecosystem Assessment

## Executive Summary

This report examines the biodiversity and ecological health of coastal temperate rainforests 
in the Pacific Northwest region, focusing on indicator species and habitat connectivity.

## Regional Biodiversity Analysis

### Mammalian Species Distribution

The region supports diverse mammalian populations across multiple elevation zones, 
with distinct patterns emerging in recent survey data.

#### Large Predators

- **Gray Wolf** (*Canis lupus*) - Recovered population in northern ranges
- **Cougar** (*Puma concolor*) - Stable throughout forested areas
- **Black Bear** (*Ursus americanus*) - Abundant across all zones

#### Medium-Sized Mammals

Population densities vary significantly by habitat type and human disturbance levels.

##### Riparian Zone Specialists
- River Otter - 2.3 individuals per km of stream
- Beaver - Colony density: 0.8 per km²
- Mink - Declining in urbanized watersheds

##### Forest Interior Species
- Pine Marten - Requires old-growth connectivity
- Fisher - Sensitive to canopy fragmentation
- Red Fox - Expanding range northward

### Avian Community Structure

#### Resident Forest Birds

Year-round residents show strong site fidelity and territory maintenance.

##### Canopy Nesters
- **Spotted Owl** (*Strix occidentalis*) - Threatened; requires >200 acres old growth
- **Marbled Murrelet** - Nests on moss-covered branches
- **Varied Thrush** - Common in dense understory

##### Cavity Nesters
Critical dependence on snag availability for breeding success.

###### Primary Cavity Excavators
- Pileated Woodpecker - Creates cavities used by 20+ species
- Hairy Woodpecker - Medium-sized cavity provider
- Downy Woodpecker - Small cavity specialist

###### Secondary Cavity Users
- Wood Duck - Requires large tree cavities near water
- Common Merganser - Competes for limited nest sites
- Northern Flying Squirrel - Nocturnal cavity occupant

## Historical Data Comparison

### Population Trends (2014-2024)

| Species | 2014 Count | 2024 Count | Change | Status |
|---------|------------|------------|--------|---------|
| Gray Wolf | 12 | 47 | +292% | Recovering |
| Fisher | 23 | 18 | -22% | Declining |
| Spotted Owl | 156 | 134 | -14% | Threatened |
| Marbled Murrelet | 2,400 | 1,850 | -23% | Endangered |

## Habitat Connectivity Analysis

### Corridor Assessment

#### Primary Wildlife Corridors

Critical linkages between protected areas maintain genetic diversity and seasonal movement patterns.

##### North-South Corridors
- **Cascade Corridor** - 45 km continuous forest
- **Olympic-Fraser Connection** - Cross-border importance
- **Coastal Riparian Network** - Salmon-dependent species pathway

##### East-West Corridors
- **Columbia River Gorge** - Migration bottleneck
- **Willamette Valley Crossings** - Highly fragmented
- **Puget Sound Lowlands** - Urban barrier zones

### Fragmentation Impact Levels

#### High Impact Zones

Areas where habitat loss exceeds 60% of historical extent.

##### Urban-Wildland Interface
- Development pressure increasing 3% annually
- Road density: >2 km/km² in affected areas
- Barrier effect on salamander dispersal

##### Agricultural Conversion Areas
- Former forest now row crops: 15,000 hectares
- Hedgerow network supports some movement
- Pesticide drift affecting amphibians

## Conservation Recommendations

### Immediate Actions Required

#### Habitat Protection
- Acquire 5,000 hectares in priority corridors
- Establish conservation easements on private lands
- Enforce stream buffer regulations

More information available at https://www.conservation.gov/habitat-protection and 
www.wildlifecoridors.org for corridor planning resources.

#### Species-Specific Interventions
- Augment Fisher population through translocation
- Protect all known Spotted Owl nesting sites
- Restore Marbled Murrelet nesting habitat

**Research Links:**
- Fisher translocation protocols: https://wildlife.org/fisher-translocation
- Spotted Owl conservation: http://www.owlconservation.org/spotted-owl
- Marbled Murrelet habitat requirements: www.seabirdhabitat.net/murrelet

### Long-Term Strategic Goals

#### Climate Adaptation Planning
Projected changes will shift suitable habitat 200-400 meters upslope by 2050.

##### Assisted Migration Considerations
- Evaluate translocation needs for low-mobility species
- Establish seed banks for foundation tree species
- Monitor for range-shifting southern species

## Monitoring Alert Types

| Alert Name | Trigger Condition | Detection Sources |
|------------|-------------------|-------------------|
| Brute Force | Multiple failed logins | SIEM, Firewall |
| SQL Injection | Malicious SQL in input | WAF, IDS |
| Data Exfiltration | Large data transfer | DLP, Network Monitor |
"""

    render_docx("/tmp/test_biodiversity.docx", test_md, title="Pacific Northwest Ecosystem Assessment")
    print("Test document created at /tmp/test_biodiversity.docx")